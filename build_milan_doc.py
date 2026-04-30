"""Generate the Milan Aegis Virtual Humanoid User Flow document.

Mirrors the structure of virtual_humanoid_user_flow_self_explanatory_template_with_button_ids.pdf
but replaces example content with Milan Aegis (read-only log analysis dashboard) workflow.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement


HEADER_BLUE = "2E5395"
HEADER_TEXT = "FFFFFF"
TITLE_BLUE = "1F3A6B"
SECTION_BLUE = "2E5395"


def set_cell_shade(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def set_cell_borders(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:color"), "888888")
        tc_borders.append(b)
    tc_pr.append(tc_borders)


def style_cell(cell, text, bold=False, color_hex=None, size=9, header=False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(str(text))
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.bold = bool(bold or header)
    if color_hex:
        run.font.color.rgb = RGBColor.from_string(color_hex)
    elif header:
        run.font.color.rgb = RGBColor.from_string(HEADER_TEXT)
    if header:
        set_cell_shade(cell, HEADER_BLUE)
    set_cell_borders(cell)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.autofit = False
    if col_widths:
        for i, w in enumerate(col_widths):
            for cell in table.columns[i].cells:
                cell.width = Inches(w)
    for i, h in enumerate(headers):
        style_cell(table.rows[0].cells[i], h, header=True, size=9)
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            style_cell(table.rows[r_idx + 1].cells[c_idx], val, size=9)
    return table


def add_title(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14 if level > 1 else 6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.bold = True
    if level == 0:
        run.font.size = Pt(24)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run.font.color.rgb = RGBColor(0, 0, 0)
    elif level == 1:
        run.font.size = Pt(15)
        run.font.color.rgb = RGBColor.from_string(SECTION_BLUE)
    else:
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor.from_string(SECTION_BLUE)
    return p


def add_para(doc, text, italic=False, bold=False, size=10, color=None, align=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(4)
    if align:
        p.alignment = align
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.italic = italic
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    return p


def add_code_block(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    return p


def main():
    doc = Document()

    for section in doc.sections:
        section.left_margin = Cm(1.2)
        section.right_margin = Cm(1.2)
        section.top_margin = Cm(1.2)
        section.bottom_margin = Cm(1.2)

    add_title(doc, "Virtual Humanoid User Flow Template", level=0)
    add_para(
        doc,
        "AI Execution Specification for the Milan Aegis Read-Only Log Analysis Dashboard",
        size=12,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    add_para(
        doc,
        "Includes detailed examples for all action types: navigation, UI/UX page maps, selectors, "
        "coordinates, pop-ups, drilldowns, dashboard, alerts, summary, audit, multi-screen behavior, "
        "validation, recovery, and humanoid narration.",
        size=10,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    add_para(
        doc,
        "How to use this document: This template is written so an AI agent or virtual humanoid can "
        "read the flow, identify the target screen, locate the UI element, perform the required "
        "action, validate the result, narrate like a human, and recover if the UI does not behave as "
        "expected.",
        italic=True, size=10, color=TITLE_BLUE,
    )

    add_table(
        doc,
        ["Document Field", "Value"],
        [
            ["Prototype application name", "Milan Aegis Log Intelligence Dashboard"],
            ["Primary user role", "SOC analyst / IT security reviewer watching the humanoid review the analyzed log baseline"],
            ["Execution actor", "Virtual humanoid operating the Milan Aegis read-only browser dashboard"],
            ["Created / regenerated on", "2026-04-30"],
            ["Template status", "Filled with Milan Aegis read-only analysis workflow"],
        ],
        col_widths=[2.4, 5.6],
    )

    add_title(doc, "1. Prototype Scenario Used Throughout This Template", level=1)
    add_para(
        doc,
        "This document uses the Milan Aegis read-only analysis prototype so every table is "
        "self-explanatory. Replace the example values with your real environment details when "
        "implementing your own flows.",
        size=10,
    )
    add_table(
        doc,
        ["Area", "Prototype Example"],
        [
            ["Web application", "Milan Aegis (read-only IT log analysis dashboard, served by milan-fe + milan-be)"],
            ["Purpose",
             "Browser dashboard that surfaces the analyzed security log baseline (KPIs, anomaly alerts, "
             "sources, policy findings, reports, audit). The system is read-only - no editing, sorting, "
             "or data mutation. The only outbound action is an email reply to milan.aegis@centific.com "
             "to Accept or Reject the analyzed findings."],
            ["Main desktop URL", "https://milan-aegis.centific.com/home"],
            ["Viewport baseline", "1440 x 900 browser viewport, 100 percent zoom"],
            ["Coordinate origin", "Top-left of browser content area after the address bar is excluded"],
            ["Primary business flow",
             "Review Milan Aegis analyzed baseline by visiting Home -> Notifications -> Dashboard -> "
             "Alerts -> Summary -> Audit, then submit Accept or Reject via the email reply Decision Panel."],
            ["Primary audience", "SOC reviewer watching the humanoid explain analyzed anomalies, KPIs, and policy findings."],
        ],
        col_widths=[2.0, 6.0],
    )

    add_title(doc, "2. Complete Flow Identification Template", level=1)
    add_table(
        doc,
        ["Field", "Required", "Description", "Filled Example"],
        [
            ["flow_id", "Yes", "Unique ID for the complete business journey.", "FLOW-MILAN-LOG-REVIEW-001"],
            ["flow_name", "Yes", "Human readable name.", "Review Milan Aegis Analyzed Log Baseline"],
            ["flow_version", "Yes", "Version used by the AI runtime.", "1.0"],
            ["parent_flow_id", "No", "Used when this flow is a child flow.", "None"],
            ["business_goal", "Yes", "Why the user wants the flow performed.",
             "Explain the analyzed log baseline and submit an Accept or Reject email decision."],
            ["persona", "Yes", "Role being simulated by humanoid.", "SOC analyst / IT security reviewer"],
            ["trigger_utterance", "Yes", "User command that starts the flow.",
             "Show me the Milan Aegis analyzed baseline and explain the anomalies."],
            ["entry_screen_id", "Yes", "First screen where execution begins.", "SCR-001-LOGIN"],
            ["expected_end_screen_id", "Yes", "Final screen after successful execution.", "SCR-010-FINAL-SUMMARY"],
            ["narration_mode", "Yes", "How much the humanoid should speak.", "Detailed guided walkthrough"],
            ["evidence_required", "Yes", "Whether screenshots/logs are required.",
             "Yes, capture screenshot at each major screen and the email-reply state"],
        ],
        col_widths=[1.5, 0.9, 2.6, 3.0],
    )

    add_title(doc, "3. Required Runtime Parameters", level=1)
    add_table(
        doc,
        ["Parameter", "Required", "Type", "Example", "How AI Uses It"],
        [
            ["session_id", "Yes", "String", "MILAN-SESSION-7781",
             "Correlates all actions, screenshots, and email-reply state."],
            ["user_role", "Yes", "String", "SOC Analyst",
             "Determines explanation depth and which screens are appropriate."],
            ["workspace_name", "Yes", "String", "Milan Aegis",
             "Confirms the active workspace branding visible in the page header pill."],
            ["dashboard_name", "Yes", "String", "Milan Aegis Dashboard",
             "Used to validate that the Dashboard route loaded the correct title."],
            ["alerts_view_id", "Yes", "String", "alerts-table",
             "Used to locate the analyzed anomaly alerts table (read-only)."],
            ["summary_view_id", "Yes", "String", "summary-screen",
             "Used to locate the Sources, Policy Findings, and Reports cards."],
            ["audit_view_id", "Yes", "String", "audit-screen",
             "Used to locate the read-only Workspace Activity log."],
            ["decision_email_address", "Yes", "Email", "milan.aegis@centific.com",
             "Destination address for the only outbound action (Accept / Reject email reply)."],
            ["target_severity", "Optional", "Enum", "critical",
             "Used to focus the humanoid on the highest-risk anomaly rows."],
            ["viewport_width", "Yes", "Number", "1440", "Used for coordinates and region map."],
            ["viewport_height", "Yes", "Number", "900", "Used for coordinates and region map."],
            ["speaker_style", "Yes", "Enum", "Professional, concise, explanatory", "Controls humanoid narration."],
        ],
        col_widths=[1.7, 0.9, 0.8, 1.7, 2.9],
    )

    add_title(doc, "4. UI/UX Screen Inventory: Where the Humanoid Goes", level=1)
    add_table(
        doc,
        ["Screen ID", "Screen / Page Name", "URL or Route", "Purpose", "Primary Regions", "Expected Visible Evidence"],
        [
            ["SCR-001-LOGIN", "Login Page", "/login",
             "Authenticate or resume session.",
             "Header, credential card, Sign In button",
             "Milan Aegis logo, username field, password field, Sign In button"],
            ["SCR-002-HOME", "Milan Aegis Home (Desktop)", "/home",
             "Main read-only landing surface with widgets and dock.",
             "Clock widget, Date widget, Notifications widget, BottomNav dock",
             "Clock time, today's date, NOTIFICATIONS count, Dashboard / Alerts / Summary / Audit dock icons"],
            ["SCR-003-NOTIFICATIONS", "Anomaly Notifications Modal", "Overlay on /home or any screen",
             "Quick preview and inspect of analyzed anomalies.",
             "Modal header, anomaly cards, Inspect button, Open Alerts Page button",
             "Title 'Anomaly Alerts', anomaly count, anomaly cards with risk score"],
            ["SCR-004-DASHBOARD", "Milan Aegis Dashboard", "/dashboard",
             "Read-only KPIs, charts, and event feed from analyzed baseline.",
             "Page header, KPI grid, charts row, event feed, dock",
             "Total Events, Cleaned Records, Duplicates Removed, Anomalies Detected tiles; Login Activity bar chart; Risk Distribution donut; Event Feed list"],
            ["SCR-005-ALERTS", "Anomaly Alerts Table", "/alerts",
             "Read-only table of analyzed anomalies with risk scores and reasons.",
             "Page header, alerts card, alerts table, status pills",
             "Columns: Timestamp, User, Event, Risk Score, Reason, Status; Analyzed badge"],
            ["SCR-006-ALERT-FOCUS", "Alert Focus State", "/alerts?focus={alert_id}",
             "Single anomaly row scrolled into view and outlined for review.",
             "Highlighted row inside alerts table",
             "Outlined alert row, focused risk score, focused status pill"],
            ["SCR-007-SUMMARY", "Analysis Summary", "/summary",
             "Read-only Sources, Policy Findings, and Reports cards.",
             "Page header, sources card, policy card, reports card",
             "Sources tiles with event counts, Policy Findings tiles, Reports tiles with status dot"],
            ["SCR-008-AUDIT", "Audit Log", "/audit",
             "Read-only trace of analysis refreshes, anomaly detection, and review activity.",
             "Page header, Workspace Activity card, activity list",
             "Activity items with type icon, title, detail, timestamp"],
            ["SCR-009-DECISION-EMAIL", "Email Reply Decision Panel", "/dashboard or /summary (embedded)",
             "Simulates the only outbound action: Accept or Reject reply to milan.aegis@centific.com.",
             "Decision card heading, description, action buttons, decision state",
             "Email reply Accept / Reject buttons, Admin records buttons, IT Admin reply state, final decision state"],
            ["SCR-010-FINAL-SUMMARY", "Humanoid Final Summary", "/home (post-review)",
             "Final spoken summary, screens visited, and recorded email decision.",
             "Summary narration, screens visited, decision recap",
             "Completed status, screens reviewed, email reply decision recorded"],
        ],
        col_widths=[1.4, 1.5, 1.4, 1.5, 1.4, 2.0],
    )

    add_title(doc, "5. Screen Layout Region Map", level=1)
    add_para(
        doc,
        "Use this to tell the AI where to look before using a selector or coordinate. Coordinates "
        "are examples based on a 1440 x 900 viewport.",
        size=10,
    )
    add_table(
        doc,
        ["Region ID", "Applies To", "Pixel Bounds x1,y1,x2,y2", "Description", "Example AI Instruction"],
        [
            ["REG-PAGE-HEADER", "All app pages", "0,0,1440,120",
             "Page header with pill, title, subtitle, and Analyzed badge.",
             "Read the pill, title, and subtitle to confirm the active screen."],
            ["REG-DOCK", "All app pages", "0,830,1440,900",
             "BottomNav dock with Dashboard, Alerts, Summary, Audit icons.",
             "Use the dock to switch between modules; check alert badge count."],
            ["REG-MAIN-CONTENT", "Most pages", "0,120,1440,830",
             "Main content area for KPIs, charts, tables, cards.",
             "Search for KPI tiles, charts, tables, and cards inside this region."],
            ["REG-HOME-WIDGETS", "SCR-002-HOME", "1000,80,1400,820",
             "Right-side stack: Date widget and Notifications widget.",
             "Open the Notifications widget here to preview anomaly alerts."],
            ["REG-HOME-CLOCK", "SCR-002-HOME", "60,160,640,500",
             "Large clock widget on the home desktop.",
             "Read time and greeting before any action on home."],
            ["REG-MODAL-OVERLAY", "SCR-003-NOTIFICATIONS", "420,140,1020,780",
             "Centered Anomaly Alerts modal overlay.",
             "Read modal title, anomaly cards, and inspect or close before any other action."],
            ["REG-KPI-GRID", "SCR-004-DASHBOARD", "24,140,1416,300",
             "Total Events, Cleaned Records, Duplicates Removed, Anomalies Detected tiles.",
             "Describe each KPI tile value and label."],
            ["REG-CHARTS-ROW", "SCR-004-DASHBOARD", "24,310,1416,600",
             "Login Activity bar chart and Risk Distribution donut.",
             "Hover bars and donut segments to read values."],
            ["REG-EVENT-FEED", "SCR-004-DASHBOARD", "24,610,1416,820",
             "Latest normalized security events list.",
             "Read severity dot, user, type, status, and timestamp per row."],
        ],
        col_widths=[1.6, 1.5, 1.5, 1.9, 2.0],
    )

    add_title(doc, "6. Element Target Registry", level=1)
    add_para(
        doc,
        "The AI should prefer stable selectors and accessible names. Coordinates are fallbacks "
        "when selectors are unavailable. Always validate visibility before action.",
        size=10,
    )

    add_title(doc, "6A. Button ID and Mapping Registry", level=2)
    add_para(
        doc,
        "Use this registry whenever a workflow depends on a button, control, or modal action. "
        "Button IDs make the action map easier to reference in templates, validations, and "
        "recovery steps.",
        size=10,
    )
    add_table(
        doc,
        ["Button ID", "Screen ID", "Button Label", "Button Mapping / Intent", "Selector / Locator", "Required Action"],
        [
            ["BTN-001", "SCR-001-LOGIN", "Sign In button", "Primary authentication action",
             'button:has-text("Sign In")', "Click"],
            ["BTN-002", "SCR-002-HOME", "Notifications widget", "Open anomaly alerts modal preview",
             "#notifications-widget-btn", "Click"],
            ["BTN-003", "SCR-003-NOTIFICATIONS", "Inspect anomaly button",
             "Drill into the focused alert row on /alerts",
             "#notif-inspect-{id}-btn", "Click"],
            ["BTN-004", "SCR-009-DECISION-EMAIL", "Email reply Accept button",
             "Submit Accept reply to milan.aegis@centific.com (only outbound action)",
             "#email-reply-accept-btn", "Click after reading findings"],
        ],
        col_widths=[1.0, 1.7, 1.6, 2.0, 1.8, 1.4],
    )

    add_table(
        doc,
        ["Element ID", "Screen ID", "Element Name", "Type", "Selector / Locator", "Fallback Location", "Required Action"],
        [
            ["EL-001", "SCR-001-LOGIN", "Username field", "Input",
             'input[name="username"], aria-label="Username"',
             "REG-MAIN-CONTENT x=610 y=340", "Type username"],
            ["EL-002", "SCR-001-LOGIN", "Password field", "Input",
             'input[name="password"], aria-label="Password"',
             "REG-MAIN-CONTENT x=610 y=405", "Type password token"],
            ["EL-003", "SCR-001-LOGIN", "Sign In button", "Button",
             'button:has-text("Sign In")',
             "REG-MAIN-CONTENT x=720 y=470", "Click"],
            ["EL-004", "SCR-002-HOME", "Notifications widget", "Card / Button",
             "#notifications-widget-btn",
             "REG-HOME-WIDGETS x=1180 y=520", "Click to open modal"],
            ["EL-005", "SCR-004-DASHBOARD", "Anomalies KPI tile", "KPI Tile",
             'data-testid="kpi-anomalies"',
             "REG-KPI-GRID x=1290 y=210", "Click for emphasis"],
            ["EL-006", "SCR-004-DASHBOARD", "Risk Distribution donut", "Chart",
             'data-testid="chart-risk-distribution"',
             "REG-CHARTS-ROW x=1080 y=440", "Hover for legend"],
            ["EL-007", "SCR-004-DASHBOARD", "Login Activity bar", "Chart bar",
             'data-testid="chart-login-activity"',
             "REG-CHARTS-ROW x=380 y=440", "Hover bar to read tooltip"],
            ["EL-008", "SCR-004-DASHBOARD", "Event feed row", "List row",
             'data-testid="event-row-{id}"',
             "REG-EVENT-FEED x=720 y=680", "Read row text"],
            ["EL-009", "SCR-005-ALERTS", "Alert table row", "Table row",
             'data-testid="alert-row-{id}"',
             "REG-ALERTS-TABLE x=720 y=320", "Read row values"],
            ["EL-010", "SCR-005-ALERTS", "Risk score badge", "Badge",
             "#alert-risk-{id}",
             "REG-ALERTS-TABLE x=900 y=320", "Read numeric score"],
            ["EL-011", "SCR-007-SUMMARY", "Sources tile", "Card tile",
             'data-testid="source-{id}"',
             "REG-MAIN-CONTENT x=300 y=260", "Read tile label, value, rationale"],
            ["EL-012", "SCR-008-AUDIT", "Workspace activity item", "List item",
             "#activity-{id}",
             "REG-MAIN-CONTENT x=720 y=320", "Read title, detail, timestamp"],
            ["EL-013", "SCR-003-NOTIFICATIONS", "Modal close button", "Button",
             "#notifications-modal-close-btn",
             "REG-MODAL-OVERLAY x=1000 y=160", "Click"],
            ["EL-014", "SCR-009-DECISION-EMAIL", "Email reply Accept button", "Button",
             "#email-reply-accept-btn",
             "REG-DECISION-PANEL x=180 y=720", "Click to submit Accept reply"],
        ],
        col_widths=[0.9, 1.6, 1.6, 1.0, 1.8, 1.6, 1.5],
    )

    add_title(doc, "7. Standard AI Action Record Template", level=1)
    add_code_block(doc, """{
  "action_id": "ACT-###",
  "target": {
    "element_id": "EL-###",
    "button_id": "BTN-###",
    "selector": "preferred stable locator",
    "button_mapping": "registry mapping or intent label",
    "fallback_region": "REG-...",
    "fallback_coordinates": {"x": 0, "y": 0}
  },
  "input_parameters": {},
  "preconditions": [],
  "execution_steps": [],
  "expected_response": {},
  "validation_rules": [],
  "humanoid_speaker_notes": {
    "before_action": "",
    "during_action": "",
    "after_action": ""
  },
  "failure_recovery": []
}""")

    add_title(doc, "8. Detailed Examples for All Action Types", level=1)
    add_table(
        doc,
        ["Action ID", "Action Type", "Category", "Screen", "Exact Example", "Parameters", "Validation", "Humanoid Speaker Note"],
        [
            ["ACT-001", "Open URL / launch session", "Browser navigation", "SCR-001-LOGIN",
             "Navigate to https://milan-aegis.centific.com/home",
             "application_url, session_id",
             "Login page or resumed Milan Aegis home is visible",
             "I will open the Milan Aegis dashboard so I can start the analyzed log review."],
            ["ACT-002", "Type text", "Form input", "SCR-001-LOGIN",
             "Enter username into EL-001", "username",
             "Username appears in field; no validation error",
             "I am entering the SOC analyst identifier for this session."],
            ["ACT-003", "Secure credential entry", "Form input", "SCR-001-LOGIN",
             "Enter password token into EL-002", "password_token",
             "Password field receives masked value",
             "I will enter the secure credential token. The value is hidden on screen."],
            ["ACT-004", "Click button", "Mouse click", "SCR-001-LOGIN",
             "Click Sign In button EL-003", "none",
             "Milan Aegis home loads within timeout",
             "I will sign in now and wait for the home desktop to load."],
            ["ACT-005", "Click dock icon", "Desktop interaction", "SCR-002-HOME",
             "Click Dashboard dock icon in BottomNav",
             "module_name=Dashboard",
             "Dashboard route loads and KPI grid is visible",
             "I will open the Milan Aegis Dashboard from the dock."],
            ["ACT-006", "Open notifications widget", "Widget open", "SCR-002-HOME",
             "Click Notifications widget EL-004",
             "none",
             "Anomaly Alerts modal opens with anomaly count",
             "I will open the notifications widget to preview the analyzed anomalies."],
            ["ACT-007", "Open page from modal", "Navigation click", "SCR-003-NOTIFICATIONS",
             "Click 'Open Alerts Page' (#notifications-open-alerts-btn)",
             "none",
             "Alerts route loads and the alerts table is visible",
             "I will open the full Alerts page to review every analyzed anomaly."],
            ["ACT-008", "Apply visual filter", "Severity scan", "SCR-005-ALERTS",
             "Visually focus rows where severity = critical (read-only; no interactive filter)",
             "target_severity",
             "Critical rows are identified by their risk-score badge color",
             "Milan is read-only, so I will visually scan the table to focus on critical anomalies."],
            ["ACT-009", "Click drilldown tile", "Drilldown", "SCR-004-DASHBOARD",
             "Click Anomalies Detected KPI tile EL-005",
             "metric_name=anomalies_detected",
             "Tile is highlighted; humanoid narrates the anomaly count and drives to /alerts",
             "I will focus on the anomalies tile to understand how many alerts the analysis surfaced."],
            ["ACT-010", "Read table row", "Data extraction", "SCR-005-ALERTS",
             "Locate the row with the highest risk score and read its fields",
             "sort_column=Risk Score, metric=highest",
             "Row id, user, event, risk score, reason captured",
             "I am reviewing the alerts and identifying the highest-risk anomaly."],
            ["ACT-011", "Hover tooltip", "Hover", "SCR-004-DASHBOARD",
             "Hover over a Login Activity bar to read its hour and event count",
             "chart_series=login_activity, x_value=hour",
             "Tooltip / title shows '{hh:00} - {n} events'",
             "I will hover over this hour bar to read the exact event count."],
            ["ACT-012", "Scroll page", "Scroll", "SCR-004-DASHBOARD",
             "Scroll down 600 px to surface the Event Feed card",
             "direction=down, distance=600",
             "Event Feed card visible",
             "I will scroll down to review the latest normalized security events."],
            ["ACT-013", "Open notifications modal", "Panel open", "SCR-002-HOME",
             "Click Notifications widget EL-004",
             "none",
             "Anomaly Alerts modal overlay appears",
             "I will open the anomaly notifications so we can preview today's alerts."],
            ["ACT-014", "Open Risk Distribution chart", "Visualization open", "SCR-004-DASHBOARD",
             "Focus the Risk Distribution donut EL-006 and its legend",
             "chart=risk_distribution",
             "Donut and Low/Medium/High/Critical legend visible",
             "I will open the risk distribution chart, which is the core visual of the dashboard."],
            ["ACT-015", "Explain Risk Distribution", "Visualization narration", "SCR-004-DASHBOARD",
             "Read each segment value: Low, Medium, High, Critical",
             "segments",
             "Humanoid explanation generated with concrete counts",
             "This donut summarizes the analyzed risk distribution across Low, Medium, High, and Critical."],
            ["ACT-016", "Advance to next chart card", "Visualization control", "SCR-004-DASHBOARD",
             "Move focus from Login Activity card to Event Feed card",
             "chart_index + 1",
             "Next card visible and described",
             "I will move to the event feed to continue the explanation."],
            ["ACT-017", "Open Summary screen", "Cross-screen open", "SCR-007-SUMMARY",
             "Click Summary dock icon (BottomNav)",
             "module_name=Summary",
             "Sources, Policy Findings, and Reports cards visible",
             "I will open the Summary screen to validate the dashboard against sources, policy, and reports."],
            ["ACT-018", "Handle pop-up", "Modal decision", "Any",
             "Read modal text and dismiss with the close button or Escape key",
             "popup_text, required_action",
             "Modal closes; underlying screen is interactive again",
             "A modal appeared. I will read it and dismiss it because Milan is read-only and no submission is required here."],
            ["ACT-019", "Switch Summary card", "Section navigation", "SCR-007-SUMMARY",
             "Move focus from Sources card to Policy Findings card to Reports card",
             "card_name",
             "Targeted card is in viewport and described",
             "I will move from sources to policy findings, and finally to reports."],
            ["ACT-020", "Read tile values", "Card review", "SCR-007-SUMMARY",
             "Read each Sources tile (label, events count, rationale)",
             "tile_set=sources",
             "Tile data captured and summarized",
             "I am reading the source health tiles to understand which systems contributed events."],
            ["ACT-021", "Re-rank by severity", "Visual re-rank", "SCR-005-ALERTS",
             "Visually scan the risk score badges (the table is read-only and not sortable)",
             "metric=Risk Score",
             "Critical badges identified; top alert noted",
             "Milan does not allow sorting, so I will visually rank by the risk score badge color."],
            ["ACT-022", "Open Audit screen", "Cross-screen open", "SCR-008-AUDIT",
             "Click Audit dock icon (BottomNav)",
             "module_name=Audit",
             "Workspace Activity card visible with activity list",
             "I will open the audit log to read the written trace of analysis activity."],
            ["ACT-023", "Search audit log", "Read scan", "SCR-008-AUDIT",
             "Visually scan activity items for type=anomaly or alert",
             "type_filter=anomaly|alert",
             "Matching activity items identified",
             "I will look for anomaly_detection entries to focus the explanation."],
            ["ACT-024", "Explain audit entry", "Section narration", "SCR-008-AUDIT",
             "Read title, detail, and timestamp of the focused activity item",
             "activity_id",
             "Activity summary generated",
             "This audit entry explains when the analysis ran and what it surfaced."],
            ["ACT-025", "Open alert reason", "Detail read", "SCR-005-ALERTS",
             "Read the full reason text in the focused alert row",
             "alert_id",
             "Reason captured into response",
             "I will read the analyzer's reason to explain why this anomaly was flagged."],
            ["ACT-026", "Increase readability", "Viewer control", "Any",
             "Increase browser zoom one step if the event feed text is small",
             "zoom_level",
             "Text becomes readable",
             "I will increase the zoom so the event feed text is comfortable to read."],
            ["ACT-027", "Switch screen via dock", "Multi-screen control", "Any",
             "Switch back to /dashboard via BottomNav Dashboard icon",
             "target_route",
             "Dashboard becomes active; KPI grid visible",
             "I will return to the dashboard to connect the audit findings to the KPIs."],
            ["ACT-028", "Close modal", "Cleanup action", "SCR-003-NOTIFICATIONS",
             "Click modal close button EL-013 or press Escape",
             "none",
             "Modal closes; underlying screen is unobstructed",
             "I will close the notifications modal now."],
            ["ACT-029", "Capture evidence", "Observability", "Any",
             "Take screenshot and store action result with screen id",
             "evidence_policy",
             "Screenshot id recorded",
             "I will capture the current screen as evidence for this step."],
            ["ACT-030", "Recover from missing element", "Error handling", "Any",
             "If element not found, fall back to direct route URL or coordinate",
             "element_id",
             "Recovery path or fail reason recorded",
             "I do not see the expected control, so I will navigate via the direct route."],
            ["ACT-031", "Submit email decision and summarize", "Outbound action + narration", "SCR-009-DECISION-EMAIL",
             "Click EL-014 (Email reply Accept) or the Reject button, then deliver final spoken summary",
             "decision in {Accept, Reject}, decision_email_address",
             "Decision state shows IT Admin reply with timestamp; final summary generated",
             "Milan is read-only; the only outbound action is the email reply. I will submit the decision and summarize what was reviewed."],
        ],
        col_widths=[0.7, 1.3, 1.0, 1.2, 1.6, 1.0, 1.4, 1.8],
    )

    add_title(doc, "9. Complete Example Flow: Milan Aegis Log Review", level=1)
    add_table(
        doc,
        ["Seq", "Action ID", "Screen From -> To", "Humanoid Does This", "UI Target / Location",
         "Expected Response", "Speaker Notes While Performing"],
        [
            ["1", "ACT-001", "Browser -> SCR-001",
             "Open Milan Aegis URL", "Address bar / application_url",
             "Login page or home loads",
             "I will open the Milan Aegis dashboard."],
            ["2", "ACT-002 to ACT-004", "SCR-001 -> SCR-002",
             "Authenticate", "Username, password, Sign In button",
             "Home desktop visible with widgets and dock",
             "I am signing in and waiting for the Milan Aegis home to load."],
            ["3", "ACT-013", "SCR-002 -> SCR-003",
             "Open notifications widget",
             "Notifications widget at right of home",
             "Anomaly Alerts modal opens",
             "I will preview the anomaly notifications first."],
            ["4", "ACT-007", "SCR-003 -> SCR-005",
             "Open the Alerts page",
             "Open Alerts Page button in modal footer",
             "Alerts table loads",
             "I will open the full alerts table."],
            ["5", "ACT-005 / ACT-027", "SCR-005 -> SCR-004",
             "Switch to Dashboard via dock",
             "Dashboard dock icon",
             "Dashboard KPIs and charts visible",
             "I will move to the dashboard to read the KPIs."],
            ["6", "ACT-009 / ACT-014 / ACT-015", "SCR-004",
             "Read KPIs and charts",
             "KPI grid, Login Activity, Risk Distribution",
             "Tile values and chart segments captured",
             "I will explain the KPIs and the risk distribution."],
            ["7", "ACT-012 / ACT-008 / ACT-010", "SCR-004 -> SCR-005",
             "Drill into highest risk anomaly",
             "Event feed row -> alerts table",
             "Highest-risk anomaly explained",
             "I will inspect the highest-risk anomaly and read its reason."],
            ["8", "ACT-017 / ACT-019 / ACT-020", "SCR-005 -> SCR-007",
             "Open Summary screen and review tiles",
             "Summary dock icon, sources/policy/reports cards",
             "Sources, Policy Findings, Reports tiles described",
             "I will validate the dashboard with sources, policy, and reports."],
            ["9", "ACT-022 / ACT-023 / ACT-024", "SCR-007 -> SCR-008",
             "Open Audit and explain activity",
             "Audit dock icon, activity list",
             "Audit items summarized",
             "I will read the analysis audit trace."],
            ["10", "ACT-027", "SCR-008 -> SCR-004",
             "Return to dashboard",
             "Dashboard dock icon",
             "Dashboard active",
             "I will return to the dashboard to connect findings."],
            ["11", "ACT-028", "Any -> SCR-002",
             "Close any open modal",
             "Modal close button or Escape",
             "Modal dismissed",
             "I will close the modal so the next action is unobstructed."],
            ["12", "ACT-029", "Any",
             "Capture evidence",
             "Screenshot policy",
             "Screenshot ids recorded",
             "I will capture screenshots of each major screen."],
            ["13", "ACT-031", "Any -> SCR-009 / SCR-010",
             "Submit email reply Accept or Reject and summarize",
             "Decision panel buttons",
             "Decision recorded; final summary delivered",
             "Milan is read-only - I will submit the email reply and deliver the final summary."],
        ],
        col_widths=[0.4, 1.3, 1.3, 1.5, 1.5, 1.3, 1.7],
    )

    add_title(doc, "10. Detailed Page Example: Milan Aegis Home", level=1)
    add_table(
        doc,
        ["Item", "Example Details"],
        [
            ["screen_id", "SCR-002-HOME"],
            ["purpose",
             "Starting surface after login. The humanoid uses the Notifications widget and the BottomNav "
             "dock (Dashboard, Alerts, Summary, Audit) to navigate the read-only analysis."],
            ["layout",
             "Top: page background. Left: large Clock widget. Right: stacked Date widget and "
             "Notifications widget. Bottom: BottomNav dock with four module icons."],
            ["primary targets",
             "Notifications widget, Dashboard dock icon, Alerts dock icon, Summary dock icon, Audit dock icon."],
            ["AI observation instruction",
             "Confirm the home loaded by detecting the Clock widget time, the Notifications widget "
             "count, and the four dock icons."],
            ["narration before action",
             "The Milan Aegis home is now open. I can see the clock, the analyzed notifications "
             "count, and the four module dock icons."],
            ["narration while acting",
             "I am opening the Notifications widget to preview today's analyzed anomalies before going to the dashboard."],
            ["success validation",
             "Notifications modal or Dashboard route is reached within 10 seconds of the action."],
            ["fallback",
             "If the dock is not visible, navigate directly to /dashboard, /alerts, /summary, or /audit. If the home itself does not load, refresh once and retry."],
        ],
        col_widths=[2.0, 6.0],
    )

    add_title(doc, "11. Detailed Page Example: Milan Aegis Dashboard", level=1)
    add_table(
        doc,
        ["UI Area", "Location", "Purpose", "Humanoid Instruction", "Validation"],
        [
            ["Page header", "REG-PAGE-HEADER",
             "Confirms the active screen ('Log Analysis' pill, title 'Dashboard').",
             "Read pill, title, and subtitle.",
             "Title text contains 'Dashboard' and pill contains 'Log Analysis'."],
            ["KPI grid", "REG-KPI-GRID",
             "Shows Total Events, Cleaned Records, Duplicates Removed, Anomalies Detected.",
             "Explain each KPI tile briefly; emphasize the Anomalies tile if it has the warn state.",
             "All four tiles have non-empty numeric values and labels."],
            ["Login Activity chart", "REG-CHARTS-ROW (left)",
             "24-hour bar chart of login events across the analyzed baseline.",
             "Hover bars to read '{hh:00} - {n} events'.",
             "Tooltip-equivalent title is readable on hover."],
            ["Risk Distribution donut", "REG-CHARTS-ROW (right)",
             "Donut of Low / Medium / High / Critical alerts.",
             "Read each segment count from the legend.",
             "Donut has segments matching legend values."],
            ["Event Feed", "REG-EVENT-FEED",
             "Latest normalized security events with severity dot, user, type, status, and time.",
             "Read severity dot color, user, event type, derived status, and timestamp.",
             "At least one event row exists, or the empty-state caption is shown."],
            ["Decision Panel (when embedded)", "REG-DECISION-PANEL",
             "Email reply controls for Accept or Reject (only outbound action).",
             "Mention the panel exists; click only when the review is complete.",
             "Email reply state and final decision state are both visible."],
        ],
        col_widths=[1.6, 1.6, 1.7, 1.8, 1.5],
    )

    add_title(doc, "12. Pop-up and Modal Examples", level=1)
    add_table(
        doc,
        ["Popup Type", "Trigger", "What AI Must Read", "Correct Action", "Example Narration", "Validation"],
        [
            ["Anomaly Notifications modal", "Clicking the home Notifications widget",
             "Title 'Anomaly Alerts', anomaly count, anomaly cards with risk score and Inspect / Open Alerts Page buttons",
             "Read cards, then click Inspect or Open Alerts Page",
             "I opened the anomaly notifications. I will inspect the highest-risk one or open the full alerts page.",
             "Modal overlay visible with at least one anomaly card or the empty-state caption."],
            ["Read-only system reminder", "Trying to edit a value",
             "Reminder that Milan Aegis is read-only and the only outbound action is an email reply",
             "Acknowledge and continue with read-only navigation",
             "Milan is read-only, so I will navigate without trying to modify any value.",
             "Reminder is dismissed; underlying screen is interactive again."],
            ["Permission / load failure", "Failure to load /alerts, /summary, /audit, or /dashboard",
             "Exact error or empty-state message",
             "Refresh once; if still failing, surface the failure and continue with available screens",
             "The screen did not load. I will refresh once and continue with the available data.",
             "Either the screen reloads or the failure is recorded with the screen id."],
            ["Session timeout", "Idle session",
             "Re-authenticate prompt or login redirect",
             "Sign in again and resume the flow at the last screen",
             "The session has timed out. I will sign in again before continuing.",
             "Home or last screen is restored."],
            ["Email reply confirmation", "Clicking Email reply Accept or Reject",
             "Decision label and confirmation message",
             "Confirm the decision because Milan's only outbound action is the email reply",
             "I am submitting the email reply because the analysis review is complete.",
             "Decision state shows IT Admin reply with timestamp."],
        ],
        col_widths=[1.4, 1.3, 1.7, 1.3, 1.7, 1.6],
    )

    add_title(doc, "13. Artifact Handling Examples", level=1)
    add_table(
        doc,
        ["Artifact Type", "Open From", "How Humanoid Reviews It", "Required Notes", "Completion Criteria"],
        [
            ["Dashboard view", "BottomNav Dashboard icon",
             "Read KPI tiles, hover Login Activity bars, read Risk Distribution donut, scan Event Feed.",
             "Mention Total Events, Anomalies Detected, dominant risk tier, and any unusual hour spike.",
             "All four KPI tiles, both charts, and the event feed are reviewed."],
            ["Alerts table", "BottomNav Alerts icon or notifications modal",
             "Read each row's timestamp, user, event, risk score, reason, and status (read-only; no sorting).",
             "Mention the highest-risk alert, its reason, and its status.",
             "The highest-risk alert is identified and explained."],
            ["Summary cards", "BottomNav Summary icon",
             "Read Sources tiles, Policy Findings tiles, and Reports tiles in order.",
             "Mention each tile's label, value, and rationale; flag any zero-event source.",
             "Sources, Policy Findings, and Reports cards are all reviewed."],
            ["Audit log", "BottomNav Audit icon",
             "Read each Workspace Activity item: type icon, title, detail, timestamp.",
             "Mention the most recent analysis run and any anomaly_detection or alert entries.",
             "The latest analysis activity items are summarized."],
            ["Notifications modal", "Home Notifications widget",
             "Read each anomaly card: risk score, title, user, rationale, severity.",
             "Mention the count and the highest-risk anomaly card.",
             "Modal contents reviewed; closed cleanly via close button or Escape."],
            ["Email decision evidence", "Decision Panel on Dashboard / Summary",
             "Read decision-state line: 'IT Admin reply: <Accept|Reject> at <timestamp>'.",
             "Cite the exact decision word and timestamp in the final summary.",
             "Decision state is captured and included in the final response object."],
        ],
        col_widths=[1.3, 1.5, 2.0, 1.8, 1.4],
    )

    add_title(doc, "14. Decision Logic and Branching Examples", level=1)
    add_table(
        doc,
        ["Decision ID", "Condition", "AI Decision", "Action", "Narration"],
        [
            ["DEC-001", "User asked to focus on critical anomalies and the alerts table is loaded",
             "Visually scan to the critical risk-score badges first (read-only - no sort)",
             "ACT-008 / ACT-021",
             "I will visually focus on critical badges first because Milan does not support sorting."],
            ["DEC-002", "Notifications widget shows count > 0",
             "Open the notifications modal before going to the dashboard",
             "ACT-013",
             "There are open anomaly notifications. I will preview them before reading the dashboard."],
            ["DEC-003", "Dashboard route loads but charts show empty state",
             "Wait and retry the dashboard fetch once",
             "Wait up to 15 seconds, then refresh /dashboard",
             "The dashboard is still loading. I will wait and refresh once before explaining it."],
            ["DEC-004", "An expected screen is missing from the dock",
             "Navigate directly via the URL route",
             "Open /alerts, /summary, /audit, or /dashboard directly",
             "The dock icon is missing, so I will navigate directly to the screen URL."],
            ["DEC-005", "User attempted to edit a field on a read-only screen",
             "Refuse the edit and explain the read-only constraint",
             "Acknowledge read-only and re-route to the email reply Decision Panel",
             "Milan is read-only. To accept or reject, I will use the email reply Decision Panel."],
            ["DEC-006", "Review is complete and all screens have been visited",
             "Submit the email reply (Accept or Reject) and produce the final summary",
             "ACT-031",
             "The review is complete. I will submit the email reply and deliver the final summary."],
        ],
        col_widths=[1.0, 2.0, 1.7, 1.7, 1.6],
    )

    add_title(doc, "15. Validation Rules and Evidence Capture", level=1)
    add_table(
        doc,
        ["Validation Type", "Rule", "Example", "Evidence to Capture"],
        [
            ["Page validation",
             "Current route must match expected route, and page header pill + title must be visible.",
             "/dashboard with title 'Dashboard' and pill 'Log Analysis'",
             "Screenshot plus detected title text"],
            ["Element validation",
             "Target element must be visible, enabled, and not covered by a modal.",
             "Anomalies Detected KPI tile visible and clickable",
             "Element ID, selector, bounding box"],
            ["Action validation",
             "Expected state change must occur after action.",
             "After clicking Notifications widget, modal overlay is visible",
             "Before/after screenshots"],
            ["Data validation",
             "Required tiles, columns, or labels must exist before summarizing.",
             "Alerts table columns Timestamp, User, Event, Risk Score, Reason, Status",
             "Captured visible table or tile values"],
            ["Narration validation",
             "Speaker notes must correspond only to the visible screen state.",
             "If KPI value is visible, mention exact value; otherwise say loading",
             "Generated narration transcript"],
            ["Recovery validation",
             "If fallback used, record the original failure and the recovery path.",
             "Missing dock icon -> direct URL -> screen reached",
             "Error log plus recovery action"],
        ],
        col_widths=[1.6, 2.4, 2.4, 1.6],
    )

    add_title(doc, "16. Error Handling Matrix", level=1)
    add_table(
        doc,
        ["Error ID", "Scenario", "Detection", "Recovery", "Final AI Response if Unresolved"],
        [
            ["ERR-001", "Page timeout", "No page header detected after 20 seconds",
             "Refresh once; if it still fails, navigate via the dock or direct URL",
             "I could not load the screen after retrying."],
            ["ERR-002", "Target element not found", "Selector and text search both fail",
             "Use fallback coordinates from REG-* mapping; then page-level scan",
             "I could not locate the expected control on this screen."],
            ["ERR-003", "Modal blocks action", "Overlay detected over main content",
             "Read modal, then close it via the close button or Escape",
             "A modal is blocking the flow and no safe close path is available."],
            ["ERR-004", "Empty data state",
             "KPI tiles read 0, alerts table empty, or audit list empty",
             "Refresh once; if still empty, narrate the empty state honestly",
             "No analyzed data is available for this batch."],
            ["ERR-005", "Read-only constraint hit",
             "Attempted edit fails or has no UI",
             "Re-route to the email reply Decision Panel; do not retry the edit",
             "Milan is read-only. The only outbound action available is the email reply."],
            ["ERR-006", "Unreadable text", "Event feed rows or tile values too small",
             "Increase browser zoom one step; reload the screen if needed",
             "The screen opened but the content is not comfortably readable."],
            ["ERR-007", "Wrong screen opened",
             "Page header pill or route does not match the expected target",
             "Use the dock or direct URL to reopen the correct screen",
             "The application opened a different screen than expected."],
            ["ERR-008", "Decision Panel state never updates",
             "Decision state remains 'not yet sent' after click",
             "Click Reset, then re-submit the email reply",
             "I could not record the email reply decision."],
        ],
        col_widths=[0.9, 1.6, 1.7, 2.1, 1.7],
    )
    add_para(
        doc,
        "This matrix can also be filled with Button ID and Button Mapping values for button-driven steps, "
        "especially when the recovery depends on a modal or dock control.",
        italic=True, size=10, color=TITLE_BLUE,
    )

    add_title(doc, "17. Humanoid Narration Script Patterns", level=1)
    add_table(
        doc,
        ["Situation", "Speaker Notes Template", "Example"],
        [
            ["Before a new screen",
             "I am going to open [screen] so we can [reason].",
             "I am going to open the Milan Aegis Dashboard so we can first see the high-level KPIs."],
            ["While loading",
             "[Screen] is loading. Once it opens, I will look for [specific items].",
             "The dashboard is loading. Once it opens, I will look for Total Events, Anomalies Detected, and the risk distribution."],
            ["After screen opens",
             "The [screen] is visible. I can see [visible regions/items].",
             "The dashboard is visible. I can see KPI tiles, the login activity chart, and the risk distribution donut."],
            ["Before drilldown",
             "I will inspect [metric] to understand [reason].",
             "I will inspect the highest-risk anomaly to understand why it was flagged."],
            ["When explaining a card / chart",
             "This [item] shows [purpose]. The key visible points are [points].",
             "This donut shows the analyzed risk distribution; Critical and High account for the largest share."],
            ["When uncertain",
             "I can see [known], but [unknown] is not clearly visible. I will [safe next step].",
             "I can see the KPI labels, but the values are still loading. I will wait one refresh."],
            ["When recovering",
             "The expected [item] is not visible, so I will try [fallback].",
             "The dock icon is not visible, so I will navigate directly to the screen URL."],
            ["Final summary (with email decision)",
             "We reviewed [screens]. The main takeaway is [insight]. I have submitted [Accept|Reject] via the email reply.",
             "We reviewed the dashboard, alerts, summary, and audit. The main takeaway is that critical anomalies cluster around late-night logins. I have submitted Accept via the email reply."],
        ],
        col_widths=[1.6, 2.7, 3.7],
    )

    add_title(doc, "18. Copy-and-Fill Blank Action Matrix", level=1)
    headers_18 = ["Seq", "Action ID", "Screen ID", "Action Type", "Target Element ID",
                  "Location / Selector", "Required Input", "Expected Response",
                  "Validation", "Speaker Notes", "Fallback"]
    table = doc.add_table(rows=4, cols=len(headers_18))
    table.style = "Table Grid"
    table.autofit = True
    for i, h in enumerate(headers_18):
        style_cell(table.rows[0].cells[i], h, header=True, size=9)
    for r in range(1, 4):
        for c in range(len(headers_18)):
            style_cell(table.rows[r].cells[c], "", size=9)

    add_title(doc, "19. Copy-and-Fill Blank Page Definition", level=1)
    blank_fields = [
        "screen_id", "screen_name", "url_or_route", "business_purpose", "entry_condition",
        "exit_condition", "layout_regions", "primary_elements", "secondary_elements",
        "popups_expected", "artifacts_available", "validation_rules",
        "speaker_notes_before", "speaker_notes_after", "fallback_navigation",
    ]
    add_table(
        doc,
        ["Field", "Value to Fill"],
        [[f, ""] for f in blank_fields],
        col_widths=[2.0, 6.0],
    )

    add_title(doc, "20. Master JSON Example for AI Runtime", level=1)
    add_code_block(doc, """{
  "flow_id": "FLOW-MILAN-LOG-REVIEW-001",
  "prototype": "Milan Aegis Log Intelligence Dashboard",
  "entry_screen_id": "SCR-001-LOGIN",
  "parameters": {
    "workspace_name": "Milan Aegis",
    "dashboard_name": "Milan Aegis Dashboard",
    "alerts_view_id": "alerts-table",
    "summary_view_id": "summary-screen",
    "audit_view_id": "audit-screen",
    "decision_email_address": "milan.aegis@centific.com",
    "target_severity": "critical"
  },
  "runtime_rules": {
    "preferred_locator_order": ["data-testid", "aria-label", "visible text", "fallback coordinates"],
    "wait_timeout_seconds": 20,
    "screenshot_policy": "capture_on_major_screen_change",
    "narration_policy": "speak_before_and_after_each_major_action",
    "do_not_infer_unreadable_content": true,
    "read_only_system": true,
    "only_outbound_action": "email_reply_to_decision_email_address"
  },
  "completion_criteria": [
    "Home opened and notifications previewed",
    "Dashboard KPIs and charts explained",
    "Highest-risk anomaly inspected on Alerts",
    "Summary sources, policy, and reports reviewed",
    "Audit activity reviewed",
    "Email reply Accept or Reject submitted",
    "Final summary delivered"
  ]
}""")

    add_title(doc, "21. Final Response Object Example", level=1)
    add_code_block(doc, """{
  "flow_id": "FLOW-MILAN-LOG-REVIEW-001",
  "status": "completed_with_email_reply_accept",
  "screens_visited": [
    "SCR-002-HOME",
    "SCR-003-NOTIFICATIONS",
    "SCR-004-DASHBOARD",
    "SCR-005-ALERTS",
    "SCR-007-SUMMARY",
    "SCR-008-AUDIT",
    "SCR-009-DECISION-EMAIL"
  ],
  "artifacts_reviewed": [
    {"type": "Dashboard view", "name": "Milan Aegis Dashboard", "result": "reviewed"},
    {"type": "Alerts table",   "name": "Anomaly Alerts",        "result": "reviewed"},
    {"type": "Summary cards",  "name": "Sources / Policy / Reports", "result": "reviewed"},
    {"type": "Audit log",      "name": "Workspace Activity",    "result": "reviewed"},
    {"type": "Notifications modal", "name": "Anomaly Alerts modal", "result": "reviewed"}
  ],
  "decision": {
    "channel": "email_reply",
    "to": "milan.aegis@centific.com",
    "value": "Accept",
    "recorded_at": "2026-04-30T10:14:22Z"
  },
  "humanoid_final_spoken_summary": "We reviewed the Milan Aegis analyzed baseline. The dashboard showed total events, cleaned records, duplicates removed, and anomalies detected; the donut highlighted a Critical / High concentration. We inspected the highest-risk alert in the Alerts table, validated the Sources, Policy Findings, and Reports cards on the Summary screen, and read the latest Workspace Activity in Audit. Because Milan is read-only, the only outbound action is the email reply - I submitted Accept to milan.aegis@centific.com.",
  "recommended_next_action": "Open renewal of detection rules for the late-night login cluster identified in the Critical band."
}""")

    out_path = r"c:\Users\User\Downloads\Log_Agent (1)\Milan_Aegis_Virtual_Humanoid_User_Flow.docx"
    doc.save(out_path)
    print(f"Wrote {out_path}")


if __name__ == "__main__":
    main()
